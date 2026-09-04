"""Ingest-time extraction of binary office documents to Markdown-ish text.

Local, deterministic, no network — called from the gateway ingest chokepoint so
docx/pdf/xlsx uploads reach the model as inline text instead of an opaque path.

Design invariants (refactoring-guidelines §6):
  §6.1  pure deterministic code — no LLM, no network.
  §6.5  fail-open: EVERY path returns None on failure, never raises into the
        request path (the caller degrades to the legacy "extract it yourself" note).
  §6.6  NEVER log document content, extracted text, or PII-bearing filenames —
        failure logs carry only the extension/format.

Library reality (2026-07-18 gateway venv): pymupdf (fitz) present; python-docx
and openpyxl MISSING → docx/xlsx run on the stdlib-zipfile fallback as the active
path; the optional-import branches auto-upgrade if the lib is ever installed.
"""
from __future__ import annotations

import logging
import os
import zipfile
from xml.etree import ElementTree as ET

logger = logging.getLogger(__name__)

MAX_EXTRACT_CHARS = 100_000  # ~100 KB cap on returned inline text
_MAX_XLSX_ROWS = 200
_MAX_XLSX_COLS = 40

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_PDF_MIME = "application/pdf"

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_S_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"


def extract_document_markdown(path: str, mime: str) -> str | None:
    """Return capped Markdown-ish text for a supported binary office document,
    or None if unsupported or extraction fails. Never raises."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx" or mime == _DOCX_MIME:
            text = _extract_docx(path)
        elif ext == ".pdf" or mime == _PDF_MIME:
            text = _extract_pdf(path)
        elif ext == ".xlsx" or mime == _XLSX_MIME:
            text = _extract_xlsx(path)
        else:
            return None
    except Exception:
        logger.debug("document_extract: extraction failed for ext=%s", ext or "?")
        return None

    if not text:
        return None
    text = text.strip()
    if not text:
        return None
    if len(text) > MAX_EXTRACT_CHARS:
        text = text[:MAX_EXTRACT_CHARS] + "\n\n[... truncated ...]"
    return text


def write_markdown_sidecar(original_path: str, markdown: str) -> str | None:
    """Persist extracted Markdown next to the cached original as <original>.md.
    Same cache dir → covered by cleanup_document_cache TTL. Returns the sidecar
    path or None on failure (never raises into the request path)."""
    try:
        sidecar_path = original_path + ".md"
        with open(sidecar_path, "w", encoding="utf-8") as f:
            f.write(markdown)
        return sidecar_path
    except Exception:
        logger.debug("document_extract: sidecar write failed")
        return None


# --- docx --------------------------------------------------------------------

def _extract_docx(path: str) -> str | None:
    try:
        import docx  # python-docx (optional; MISSING in gateway venv today)
    except ImportError:
        return _extract_docx_stdlib(path)
    document = docx.Document(path)
    lines: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _extract_docx_stdlib(path: str) -> str | None:
    with zipfile.ZipFile(path) as zf:
        with zf.open("word/document.xml") as f:
            root = ET.parse(f).getroot()
    lines: list[str] = []
    for para in root.iter(f"{_W_NS}p"):
        line = "".join(node.text or "" for node in para.iter(f"{_W_NS}t"))
        if line.strip():
            lines.append(line)
    return "\n".join(lines)


# --- pdf / xlsx placeholders (implemented in Tasks 4 & 5) --------------------

def _extract_pdf(path: str) -> str | None:
    import fitz  # pymupdf — text layer only, no OCR (non-goal)
    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            txt = page.get_text("text")
            if txt.strip():
                parts.append(txt)
    return "\n\n".join(parts)


def _extract_xlsx(path: str) -> str | None:
    try:
        import openpyxl  # optional (MISSING in gateway venv today)
    except ImportError:
        return _extract_xlsx_stdlib(path)
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    out: list[str] = []
    try:
        for ws in wb.worksheets:
            out.append(f"## {ws.title}")
            for r, row in enumerate(ws.iter_rows(values_only=True)):
                if r >= _MAX_XLSX_ROWS:
                    break
                cells = ["" if v is None else str(v) for v in row[:_MAX_XLSX_COLS]]
                out.append("| " + " | ".join(cells) + " |")
    finally:
        wb.close()
    return "\n".join(out)


def _extract_xlsx_stdlib(path: str) -> str | None:
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        shared: list[str] = []
        if "xl/sharedStrings.xml" in names:
            with zf.open("xl/sharedStrings.xml") as f:
                sroot = ET.parse(f).getroot()
            for si in sroot.iter(f"{_S_NS}si"):
                shared.append("".join(t.text or "" for t in si.iter(f"{_S_NS}t")))
        sheet_files = sorted(
            n for n in names
            if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")
        )
        out: list[str] = []
        for idx, name in enumerate(sheet_files, 1):
            out.append(f"## Sheet{idx}")
            with zf.open(name) as f:
                root = ET.parse(f).getroot()
            for r, row in enumerate(root.iter(f"{_S_NS}row")):
                if r >= _MAX_XLSX_ROWS:
                    break
                cells = [
                    _xlsx_cell_text(c, shared) for c in list(row)[:_MAX_XLSX_COLS]
                ]
                out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


def _xlsx_cell_text(c, shared: list[str]) -> str:
    t = c.get("t")
    if t == "s":  # shared-string index
        v = c.find(f"{_S_NS}v")
        if v is not None and v.text and v.text.isdigit():
            i = int(v.text)
            if 0 <= i < len(shared):
                return shared[i]
        return ""
    if t == "inlineStr":
        is_el = c.find(f"{_S_NS}is")
        if is_el is not None:
            return "".join(x.text or "" for x in is_el.iter(f"{_S_NS}t"))
        return ""
    v = c.find(f"{_S_NS}v")
    return v.text if (v is not None and v.text) else ""
